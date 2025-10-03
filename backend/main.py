from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any, Literal
import httpx
import requests
import subprocess
import json
import os
from dotenv import load_dotenv
import asyncio
import aiofiles
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches
import markdown
# PDF generation temporarily disabled for deployment
import re
import db  # 导入数据库模块
import base64
from io import BytesIO

# 加载环境变量
load_dotenv()

app = FastAPI(title="Indus AI Dialogue Forge API", version="1.0.0")

# API重试配置
API_RETRY_COUNT = int(os.getenv("API_RETRY_COUNT", "3"))
API_TIMEOUT = int(os.getenv("API_TIMEOUT", "60"))

# 应用启动时初始化数据库
@app.on_event("startup")
async def startup_event():
    await db.init_db()
    print("数据库初始化完成")

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    # 开发环境放宽跨域，前端端口可能变化（8081/8082/8083/8084...）
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 数据模型
class Message(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    messages: List[Message]
    model: str = "deepseek-chat"
    temperature: float = 0.7
    max_tokens: int = 4000
    stream: bool = True
    output_format: str = "text"  # text, pdf, docx, markdown
    conversation_id: Optional[str] = None  # 本地数据库对话ID
    dify_conversation_id: Optional[str] = None  # Dify 会话ID，用于与 Dify 持续对话
    files: Optional[List[Dict[str, Any]]] = None  # 文件列表
    workflow_type: str = "prd"  # prd, prompt - 指定使用哪个工作流
    page_data: Optional[Dict[str, Any]] = None  # 页面数据（用于prompt生成）

class ChatResponse(BaseModel):
    id: str
    object: str
    created: int
    model: str
    choices: List[Dict[str, Any]]
    usage: Dict[str, int]

# Dify API 配置（可通过环境变量覆盖），默认指向教研环境网关
# 如需切换，请在运行环境中设置 DIFY_API_BASE 与相应的 API_KEY
DIFY_API_BASE = os.getenv("DIFY_API_BASE", "http://teach.excelmaster.ai/v1")

# PRD生成工作流配置
DIFY_PRD_API_KEY = os.getenv("DIFY_PRD_API_KEY", "app-wiFSsheVuALpQ5cN7LrPv5Lb")

# Prompt生成工作流配置  
DIFY_PROMPT_API_KEY = os.getenv("DIFY_PROMPT_API_KEY", "app-6tFKntYIPDWzq2toScD1XIiY")

# OpenRouter API 配置
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"

# 优先调用工作流接口（/workflows/run）。如需改为应用聊天接口（/chat-messages），将该值设为 "chat"
DIFY_API_CHANNEL = os.getenv("DIFY_API_CHANNEL", "workflow")  # workflow | chat

def _dify_endpoint() -> str:
    if DIFY_API_CHANNEL.lower() == "chat":
        return f"{DIFY_API_BASE}/chat-messages"
    return f"{DIFY_API_BASE}/workflows/run"

if not DIFY_PRD_API_KEY:
    print("警告: 未设置DIFY_PRD_API_KEY环境变量")
if not DIFY_PROMPT_API_KEY:
    print("警告: 未设置DIFY_PROMPT_API_KEY环境变量")
if not OPENROUTER_API_KEY:
    print("警告: 未设置OPENROUTER_API_KEY环境变量")
else:
    print(f"✅ OpenRouter API Key已加载: {OPENROUTER_API_KEY[:15]}...")

# API错误处理函数
async def handle_dify_api_error(response: httpx.Response, context: str = "") -> str:
    """
    处理Dify API错误并返回用户友好的错误信息
    """
    try:
        error_data = response.json()
        error_message = error_data.get('error', {}).get('message', '')
        
        if response.status_code == 403:
            if 'quota' in error_message.lower():
                return f"API配额不足，请联系管理员充值或更换API密钥。{context}"
            else:
                return f"API访问被拒绝，请检查API密钥是否正确。{context}"
        elif response.status_code == 429:
            return f"API请求频率过高，请稍后再试。{context}"
        elif response.status_code == 401:
            return f"API密钥无效，请检查配置。{context}"
        else:
            return f"API错误 ({response.status_code}): {error_message} {context}"
    except:
        return f"API错误 ({response.status_code}): {await response.aread().decode(errors='ignore')} {context}"

# API重试函数
async def retry_api_call(func, max_retries: int = API_RETRY_COUNT, delay: float = 1.0):
    """
    带重试机制的API调用
    """
    for attempt in range(max_retries):
        try:
            return await func()
        except httpx.HTTPStatusError as e:
            if e.response.status_code == 403 and 'quota' in str(e.response.text):
                # 配额不足不重试
                raise
            if attempt == max_retries - 1:
                raise
            print(f"API调用失败 (尝试 {attempt + 1}/{max_retries}): {e}")
            await asyncio.sleep(delay * (2 ** attempt))  # 指数退避
        except Exception as e:
            if attempt == max_retries - 1:
                raise
            print(f"API调用失败 (尝试 {attempt + 1}/{max_retries}): {e}")
            await asyncio.sleep(delay)

# 文件存储目录
UPLOAD_DIR = "uploads"
GENERATED_DIR = "generated"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)

# 辅助函数
async def upload_file_to_dify(file_data: bytes, filename: str, mime_type: str) -> Optional[str]:
    """
    上传文件到Dify API并返回upload_file_id
    按照经验总结的两步上传流程：先上传文件获取ID，然后在聊天中引用
    """
    try:
        print(f"开始上传文件到Dify: {filename}, 类型: {mime_type}")
        
        async with httpx.AsyncClient(timeout=60.0) as client:
            # 构造multipart/form-data格式的请求
            files = {
                'file': (filename, BytesIO(file_data), mime_type)
            }
            data = {
                'user': 'abc-123'  # 必需的用户标识
            }
            headers = {
                "Authorization": f"Bearer {DIFY_PRD_API_KEY}"  # 默认使用PRD API Key
            }
            
            # 调用Dify文件上传端点，处理可能的重定向问题
            upload_urls = [
                f"{DIFY_API_BASE}/files/upload",
                f"{DIFY_API_BASE}/files/upload/",
            ]
            
            for upload_url in upload_urls:
                try:
                    print(f"尝试上传到: {upload_url}")
                    response = await client.post(
                        upload_url,
                        files=files,
                        data=data,
                        headers=headers,
                        follow_redirects=True
                    )
                    
                    print(f"Dify文件上传响应: {response.status_code}")
                    print(f"响应内容: {response.text}")
                    
                    # 接受201和200状态码
                    if response.status_code in [200, 201]:
                        result = response.json()
                        upload_file_id = result.get('id')
                        if upload_file_id:
                            print(f"文件上传成功，获得upload_file_id: {upload_file_id}")
                            return upload_file_id
                        else:
                            print(f"文件上传响应中缺少id字段: {result}")
                    else:
                        print(f"上传URL {upload_url} 失败: {response.status_code} - {response.text}")
                        continue  # 尝试下一个URL
                
                except Exception as url_error:
                    print(f"上传URL {upload_url} 出错: {str(url_error)}")
                    continue  # 尝试下一个URL
            
            print("所有上传URL都失败")
            return None
                
    except Exception as e:
        print(f"上传文件到Dify时出错: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        return None

def is_image_file(filename: str, mime_type: str) -> bool:
    """检查文件是否为图片"""
    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp']
    image_mimes = ['image/jpeg', 'image/png', 'image/gif', 'image/webp', 'image/bmp']
    
    file_ext = os.path.splitext(filename.lower())[1]
    return file_ext in image_extensions or mime_type in image_mimes

async def convert_file_to_dify_format(file_info: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """
    将前端文件信息转换为Dify API格式
    按照经验：inputs.input_img必须是文件对象数组，不是单个文件ID
    """
    try:
        file_type = file_info.get('type')
        file_url = file_info.get('url')
        filename = file_info.get('name')
        mime_type = file_info.get('mime_type', file_info.get('type', 'application/octet-stream'))
        
        # 如果是base64格式的图片
        if file_url and file_url.startswith('data:'):
            # 解析data URL
            header, data = file_url.split(',', 1)
            mime_type = header.split(';')[0].split(':')[1]
            file_data = base64.b64decode(data)
            
            # 上传到Dify获取文件ID
            upload_file_id = await upload_file_to_dify(file_data, filename, mime_type)
            
            if upload_file_id:
                # 按照经验构造正确的格式
                return {
                    "type": "image" if is_image_file(filename, mime_type) else "file",
                    "transfer_method": "local_file",
                    "upload_file_id": upload_file_id
                }
        
        return None
        
    except Exception as e:
        print(f"转换文件格式时出错: {str(e)}")
        return None

# 文件生成函数
async def generate_markdown_file(content: str, filename: str) -> str:
    """生成Markdown文件"""
    file_path = os.path.join(GENERATED_DIR, filename)
    async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
        await f.write(content)
    return file_path

async def generate_docx_file(content: str, filename: str) -> str:
    """生成DOCX文件"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('AI生成文档', 0)
    title.alignment = 1  # 居中对齐
    
    # 添加时间戳
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()  # 空行
    
    # 处理内容，支持基本的Markdown格式
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # 标题处理
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            # 粗体
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('```'):
            # 代码块开始/结束，忽略
            continue
        else:
            # 普通段落
            doc.add_paragraph(line)
    
    file_path = os.path.join(GENERATED_DIR, filename)
    doc.save(file_path)
    return file_path

async def generate_pdf_file(content: str, filename: str) -> str:
    """生成PDF文件"""
    file_path = os.path.join(GENERATED_DIR, filename)
    
    # PDF generation temporarily disabled for deployment
    raise HTTPException(status_code=501, detail="PDF generation temporarily disabled for deployment")

async def generate_file(content: str, output_format: str) -> Optional[Dict[str, str]]:
    """根据格式生成文件"""
    if output_format == "text":
        return None
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_id = str(uuid.uuid4())[:8]
    
    try:
        if output_format == "markdown":
            filename = f"ai_response_{timestamp}_{file_id}.md"
            file_path = await generate_markdown_file(content, filename)
            return {
                "filename": filename,
                "url": f"/api/files/{filename}",
                "mime_type": "text/markdown"
            }
        elif output_format == "docx":
            filename = f"ai_response_{timestamp}_{file_id}.docx"
            file_path = await generate_docx_file(content, filename)
            return {
                "filename": filename,
                "url": f"/api/files/{filename}",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        elif output_format == "pdf":
            filename = f"ai_response_{timestamp}_{file_id}.pdf"
            file_path = await generate_pdf_file(content, filename)
            return {
                "filename": filename,
                "url": f"/api/files/{filename}",
                "mime_type": "application/pdf"
            }
    except Exception as e:
        print(f"文件生成错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        return None
    
    return None

@app.get("/")
async def root():
    return {"message": "Indus AI Dialogue Forge API", "status": "running"}

@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

# 新的简化prompt生成接口，用于Chrome插件调用
class PromptGenerateRequest(BaseModel):
    prompt: str
    screenshot: Optional[str] = None

@app.post("/api/generate-prompt")
async def generate_prompt_simple(request: PromptGenerateRequest):
    """
    简化的prompt生成接口，专为Chrome插件调用
    接收prompt文本和可选的截图，返回生成的内容
    """
    try:
        print(f"收到Chrome插件prompt生成请求, prompt长度: {len(request.prompt)}")
        
        # 检查API密钥是否配置
        if not OPENROUTER_API_KEY:
            raise HTTPException(status_code=500, detail="OpenRouter API密钥未配置")
        
        print(f"使用API密钥: {OPENROUTER_API_KEY[:15]}...")
        
        # 验证API密钥是否有效
        if not OPENROUTER_API_KEY.startswith('sk-or-v1-') or len(OPENROUTER_API_KEY) < 50:
            raise HTTPException(status_code=500, detail="OpenRouter API密钥格式不正确，请检查并更新")
        
        # 构建OpenRouter请求
        messages = [
            {"role": "system", "content": "你是一个专业的编程助手，专门帮助开发者分析网页元素并生成详细的编程实现指令。"},
            {"role": "user", "content": request.prompt}
        ]
        
        # 如果有截图，添加到消息中
        if request.screenshot:
            user_content = [
                {"type": "text", "text": request.prompt},
                {"type": "image_url", "image_url": {"url": request.screenshot, "detail": "high"}}
            ]
            messages[1]["content"] = user_content
        
        openrouter_payload = {
            "model": "openai/gpt-4o-mini",
            "messages": messages,
            "temperature": 0.7,
            "max_tokens": 4000,
            "stream": False
        }
        
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://chrome-extension-ai-assistant.com",
            "X-Title": "AI Programming Assistant"
        }
        
        print(f"发送请求到OpenRouter API...")
        
        # 使用requests同步调用作为备用方案
        try:
            import requests
            print(f"使用requests库进行API调用...")
            
            response = requests.post(
                f"{OPENROUTER_BASE_URL}/chat/completions",
                json=openrouter_payload,
                headers=headers,
                timeout=60,
                verify=False
            )
            
            if response.status_code != 200:
                error_text = response.text
                print(f"OpenRouter API错误详情: {error_text}")
                raise HTTPException(status_code=response.status_code, 
                                  detail=f"OpenRouter API错误 (status={response.status_code}): {error_text}")
            
            openrouter_data = response.json()
            print(f"OpenRouter响应成功，数据长度: {len(str(openrouter_data))}")
                
        except requests.exceptions.Timeout:
            raise HTTPException(status_code=408, detail="请求超时")
        except requests.exceptions.RequestException as e:
            print(f"网络请求错误: {str(e)}")
            raise HTTPException(status_code=500, detail=f"网络请求错误: {str(e)}")
        except Exception as e:
            print(f"API请求异常: {str(e)}")
            raise HTTPException(status_code=500, detail=f"API请求异常: {str(e)}")
        
        # 提取AI回复内容
        ai_content = ""
        if 'choices' in openrouter_data and len(openrouter_data['choices']) > 0:
            ai_content = openrouter_data['choices'][0].get('message', {}).get('content', '')
        
        if not ai_content:
            raise HTTPException(status_code=500, detail="AI生成内容为空")
        
        print(f"成功生成prompt，长度: {len(ai_content)}")
        
        return {
            "prompt": ai_content,
            "content": ai_content,  # 兼容不同的前端字段名
            "success": True
        }
            
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="请求超时")
    except httpx.RequestError as e:
        print(f"网络请求错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"网络请求错误: {str(e)}")
    except Exception as e:
        print(f"生成prompt错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {str(e)}")

@app.post("/api/prompt")
async def generate_prompt_with_openrouter(request: ChatRequest):
    """
    使用OpenRouter API生成编程prompt，专门用于智能Prompt功能
    """
    try:
        print(f"开始处理OpenRouter prompt生成请求...")
        
        # 从请求中提取用户消息和页面数据
        user_messages = [msg for msg in request.messages if msg.role == "user"]
        if not user_messages:
            raise HTTPException(status_code=400, detail="请求中必须包含用户消息")
        
        user_message = user_messages[-1]
        page_data = request.page_data
        
        if not page_data:
            raise HTTPException(status_code=400, detail="缺少页面数据")
        
        # 构建详细的prompt
        element_data = page_data.get("element", {})
        page_context = page_data.get("pageContext", {})
        
        # 构建HTML结构描述
        element_html = f"<{element_data.get('tagName', 'div')}"
        if element_data.get('attributes'):
            for attr, value in element_data.get('attributes', {}).items():
                element_html += f' {attr}="{value}"'
        element_html += f">{element_data.get('directText', '')}</{element_data.get('tagName', 'div')}>"
        
        # 构建CSS样式描述  
        element_css = ""
        if element_data.get('styles'):
            important_props = ['display', 'position', 'width', 'height', 'background-color', 
                             'color', 'font-size', 'font-family', 'border', 'border-radius', 
                             'padding', 'margin', 'flex-direction', 'justify-content', 'align-items']
            
            for prop, value in element_data.get('styles', {}).items():
                if prop in important_props:
                    element_css += f"{prop}: {value}; "
        
        # 构建详细的系统prompt
        system_prompt = """你是一个专业的编程助手，专门帮助开发者分析网页元素并生成详细的编程实现指令。

请根据提供的网页元素信息，生成包含以下内容的专业编程指令：

1. **HTML结构分析**：分析元素的HTML结构，说明如何实现
2. **CSS样式实现**：提供完整的CSS代码，包括布局、样式、响应式设计
3. **JavaScript交互**：如果有交互功能，提供JavaScript实现代码
4. **响应式设计**：提供移动端适配建议
5. **可访问性优化**：提供无障碍访问优化建议
6. **最佳实践**：提供代码优化和性能建议

请确保生成的指令详细、准确、可执行，适合开发者直接使用。"""

        # 构建用户查询
        user_query = f"""请分析以下网页元素并生成详细的编程实现指令：

**元素基本信息**
- 标签类型: {element_data.get('tagName', 'div')}
- CSS类名: {element_data.get('attributes', {}).get('class', '无')}
- ID: {element_data.get('attributes', {}).get('id', '无')}
- 页面来源: {page_context.get('domain', '')}

**HTML代码**
{element_html}

**重要CSS样式**
{element_css}

**元素属性**"""

        # 添加其他重要属性
        attrs = element_data.get('attributes', {})
        for attr, value in attrs.items():
            if attr not in ['class', 'id']:
                user_query += f"\n- {attr}: {value}"
        
        # 添加尺寸信息
        dimensions = element_data.get('dimensions', {})
        if dimensions:
            user_query += f"""

**元素尺寸**
- 宽度: {dimensions.get('width', '未知')}px
- 高度: {dimensions.get('height', '未知')}px"""

        user_query += """

请生成包含HTML结构、CSS样式、JavaScript交互（如有）、响应式设计和可访问性优化的完整编程实现指令。"""

        # 构建OpenRouter请求
        openrouter_payload = {
            "model": "openai/gpt-4o-mini",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_query}
            ],
            "temperature": 0.7,
            "max_tokens": 4000,
            "stream": request.stream
        }
        
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://chrome-extension-ai-assistant.com",
            "X-Title": "AI Programming Assistant"
        }
        
        print(f"发送请求到OpenRouter: {openrouter_payload['model']}")
        print(f"查询长度: {len(user_query)} 字符")
        
        if request.stream:
            # 流式响应
            async def generate():
                full_content = ""
                
                try:
                    async with httpx.AsyncClient(timeout=60.0) as client:
                        async with client.stream(
                            "POST",
                            f"{OPENROUTER_BASE_URL}/chat/completions",
                            json=openrouter_payload,
                            headers={**headers, "Accept": "text/event-stream"},
                        ) as response:
                            if response.status_code != 200:
                                error_text = (await response.aread()).decode(errors="ignore")
                                error_msg = f"OpenRouter API错误 (status={response.status_code}): {error_text}"
                                print(f"API错误: {error_msg}")
                                yield f"data: {json.dumps({'error': {'message': error_msg}}, ensure_ascii=False)}\n\n"
                                yield "data: [DONE]\n\n"
                                return

                            print("开始接收 OpenRouter 流式数据...")
                            buffer = ""
                            async for chunk in response.aiter_text():
                                if chunk:
                                    buffer += chunk
                                    # 处理完整的行
                                    while '\n' in buffer:
                                        line, buffer = buffer.split('\n', 1)
                                        if line.strip():
                                            if line.startswith('data: '):
                                                try:
                                                    data = line[6:].strip()
                                                    if not data or data == '[DONE]':
                                                        pass
                                                    else:
                                                        parsed = json.loads(data)
                                                        if 'choices' in parsed and len(parsed['choices']) > 0:
                                                            delta = parsed['choices'][0].get('delta', {})
                                                            if 'content' in delta:
                                                                delta_text = delta['content']
                                                                full_content += delta_text
                                                                transformed = {
                                                                    "choices": [
                                                                        {"delta": {"content": delta_text}}
                                                                    ]
                                                                }
                                                                yield f"data: {json.dumps(transformed, ensure_ascii=False)}\n\n"
                                                except Exception as e:
                                                    print(f"解析OpenRouter流式数据行失败: {e}")
                                                    pass
                            
                            # 处理剩余的buffer
                            if buffer.strip():
                                try:
                                    if buffer.startswith('data: '):
                                        data = buffer[6:].strip()
                                        if data and data != '[DONE]':
                                            parsed = json.loads(data)
                                            if 'choices' in parsed and len(parsed['choices']) > 0:
                                                delta = parsed['choices'][0].get('delta', {})
                                                if 'content' in delta:
                                                    delta_text = delta['content']
                                                    full_content += delta_text
                                                    transformed = {"choices": [{"delta": {"content": delta_text}}]}
                                                    yield f"data: {json.dumps(transformed, ensure_ascii=False)}\n\n"
                                except Exception:
                                    pass
                            
                            # 提供 HTML 渲染
                            try:
                                if full_content and isinstance(full_content, str):
                                    html_output = markdown.markdown(full_content, extensions=['extra', 'sane_lists'])
                                    html_event = {"type": "html", "html": html_output}
                                    yield f"data: {json.dumps(html_event, ensure_ascii=False)}\n\n"
                            except Exception as _e:
                                print(f"HTML 渲染失败: {_e}")
                            
                            yield "data: [DONE]\n\n"
                            
                except Exception as e:
                    print(f"OpenRouter流式处理错误: {repr(e)}")
                    import traceback
                    print(traceback.format_exc())
                    error_response = {
                        "error": {
                            "message": f"OpenRouter流式处理错误: {repr(e)}"
                        }
                    }
                    yield f"data: {json.dumps(error_response)}\n\n"
                    yield "data: [DONE]\n\n"
            
            return StreamingResponse(
                generate(),
                media_type="text/event-stream",
                headers={
                    "Cache-Control": "no-cache, no-transform",
                    "Pragma": "no-cache",
                    "Connection": "keep-alive",
                    "X-Accel-Buffering": "no",
                }
            )
        
        else:
            # 阻塞响应
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    f"{OPENROUTER_BASE_URL}/chat/completions",
                    json=openrouter_payload,
                    headers=headers,
                )
                
                if response.status_code != 200:
                    error_text = await response.aread()
                    error_text = error_text.decode('utf-8')
                    print(f"OpenRouter API错误详情: {error_text}")
                    raise HTTPException(status_code=response.status_code, 
                                      detail=f"OpenRouter API错误 (status={response.status_code}): {error_text}")
                
                openrouter_data = response.json()
                print(f"OpenRouter响应数据: {json.dumps(openrouter_data, ensure_ascii=False, indent=2)}")

                # 提取AI回复内容
                ai_content = ""
                if 'choices' in openrouter_data and len(openrouter_data['choices']) > 0:
                    ai_content = openrouter_data['choices'][0].get('message', {}).get('content', '')

                # 构造兼容的返回体
                compatible = {
                    "id": openrouter_data.get('id', str(uuid.uuid4())),
                    "object": "chat.completion",
                    "created": int(datetime.now().timestamp()),
                    "model": "openai/gpt-4o-mini",
                    "choices": [
                        {
                            "index": 0,
                            "message": {"role": "assistant", "content": ai_content},
                            "finish_reason": "stop",
                        }
                    ],
                    "usage": openrouter_data.get('usage', {
                        "prompt_tokens": 0,
                        "completion_tokens": 0,
                        "total_tokens": 0,
                    }),
                }

                # 提供 HTML 渲染字段
                try:
                    if ai_content and isinstance(ai_content, str):
                        compatible['html'] = markdown.markdown(ai_content, extensions=['extra', 'sane_lists'])
                except Exception as _e:
                    print(f"阻塞模式 HTML 渲染失败: {_e}")

                return compatible
                
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="请求超时")
    except httpx.RequestError as e:
        print(f"网络请求错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"网络请求错误: {str(e)}")
    except Exception as e:
        print(f"服务器内部错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {str(e)}")

@app.post("/api/chat")
async def chat_with_dify(request: ChatRequest):
    """
    通过 Dify 工作流进行对话，支持PRD生成和Prompt生成两个工作流。
    - 流式：将 Dify SSE 事件转换为 OpenAI/DeepSeek 风格的 delta 流
    - 阻塞：将 Dify 的 answer 映射到 choices[0].message.content
    """
    
    # 根据workflow_type选择对应的API Key
    if request.workflow_type == "prompt":
        api_key = DIFY_PROMPT_API_KEY
        if not api_key:
            raise HTTPException(status_code=500, detail="Prompt生成工作流API密钥未配置")
    else:  # 默认使用PRD工作流
        api_key = DIFY_PRD_API_KEY
        if not api_key:
            raise HTTPException(status_code=500, detail="PRD生成工作流API密钥未配置")
    
    try:
        print(f"开始处理聊天请求，工作流类型: {request.workflow_type}, API Key: {api_key[:8]}...")
        
        # 从请求中提取用户消息（取最后一个，即当前轮的输入）
        user_messages = [msg for msg in request.messages if msg.role == "user"]
        if not user_messages:
            raise HTTPException(status_code=400, detail="请求中必须包含用户消息")
        
        user_message = user_messages[-1]  # 取最后一个用户消息
        print(f"提取到的用户消息: {user_message.content[:100]}...")
        print(f"总共有 {len(user_messages)} 条用户消息，使用最后一条")
        
        # 创建或获取对话ID
        conversation_id = request.conversation_id
        
        if not conversation_id:
            # 创建新对话
            title = user_message.content[:30] + "..." if len(user_message.content) > 30 else user_message.content
            conversation_id = await db.create_conversation(title=title, model=request.model)
            print(f"创建新对话: {conversation_id}")
        else:
            # 验证对话是否存在
            conversation = await db.get_conversation(conversation_id)
            if not conversation:
                raise HTTPException(status_code=404, detail="指定的对话不存在")
            print(f"使用现有对话: {conversation_id}")
        
        # 存储用户消息
        user_message_id = await db.add_message(
            conversation_id=conversation_id,
            role="user",
            content=user_message.content
        )
        print(f"添加用户消息: {user_message_id}")
        
        # 处理文件上传（如果有）
        dify_files = []
        if request.files:
            print(f"处理 {len(request.files)} 个文件")
            for file_info in request.files:
                converted_file = await convert_file_to_dify_format(file_info)
                if converted_file:
                    dify_files.append(converted_file)
                    print(f"文件转换成功: {converted_file}")
                else:
                    print(f"文件转换失败: {file_info}")
        
        # 构建Dify请求数据
        user_query = user_message.content
        
        # 基础payload
        dify_payload = {
            "inputs": {"query": user_query, "text": user_query},
            "query": user_query,  # 若对接到 chat-messages 也可工作
            "response_mode": "streaming" if request.stream else "blocking",
            "conversation_id": (request.dify_conversation_id or ""),
            "user": "abc-123",
        }
        
        # 根据工作流类型添加特定输入
        if request.workflow_type == "prompt" and request.page_data:
            # Prompt生成工作流需要页面数据
            page_data = request.page_data
            element_data = page_data.get("element", {})
            
            # 构建HTML结构描述
            element_html = f"<{element_data.get('tagName', 'div')}"
            if element_data.get('attributes'):
                for attr, value in element_data.get('attributes', {}).items():
                    element_html += f' {attr}="{value}"'
            element_html += f">{element_data.get('directText', '')}</{element_data.get('tagName', 'div')}>"
            
            # 构建CSS样式描述  
            element_css = ""
            if element_data.get('styles'):
                # 只选择最重要的CSS属性，减少长度
                important_props = ['display', 'position', 'width', 'height', 'background-color', 
                                 'color', 'font-size', 'font-family', 'border', 'border-radius', 
                                 'padding', 'margin', 'flex-direction', 'justify-content', 'align-items']
                
                for prop, value in element_data.get('styles', {}).items():
                    if prop in important_props:
                        element_css += f"{prop}: {value}; "
            
            # 构建详细的分析查询，包含完整元素信息
            detailed_query = f"""请基于以下网页元素信息生成详细的编程实现指令：

**元素基本信息**
- 标签类型: {element_data.get('tagName', 'div')}
- CSS类名: {element_data.get('attributes', {}).get('class', '无')}
- ID: {element_data.get('attributes', {}).get('id', '无')}
- 页面来源: {page_data.get("pageContext", {}).get("domain", "")}

**HTML代码**
{element_html}

**重要CSS样式**
{element_css}

**元素属性**"""

            # 添加其他重要属性
            attrs = element_data.get('attributes', {})
            for attr, value in attrs.items():
                if attr not in ['class', 'id']:
                    detailed_query += f"\n- {attr}: {value}"
            
            # 添加尺寸信息
            dimensions = element_data.get('dimensions', {})
            if dimensions:
                detailed_query += f"""

**元素尺寸**
- 宽度: {dimensions.get('width', '未知')}px
- 高度: {dimensions.get('height', '未知')}px"""

            detailed_query += """

**请生成包含以下内容的编程指令：**
1. HTML结构分析和实现方法
2. CSS样式详细说明和代码
3. 如果有交互功能，请说明JavaScript实现
4. 响应式设计建议
5. 可访问性优化建议

请生成详细、专业的编程实现指令。"""
            
            # 为了符合Dify限制，使用简化版本作为实际传递的参数
            dify_payload["inputs"].update({
                "query": detailed_query[:2000],  # 限制长度避免超限
                "text": detailed_query[:2000],
                "element_type": element_data.get('tagName', 'div')[:20],
                "page_domain": page_data.get("pageContext", {}).get("domain", "")[:30],
            })
            
            print(f"添加页面数据到Prompt工作流: {page_data.get('pageContext', {}).get('domain', 'unknown')}")
            print(f"元素类型: {element_data.get('tagName', 'div')}")
            print(f"查询长度: {len(detailed_query)} 字符")
            print(f"传递给Dify的查询: {detailed_query[:200]}...")
        
        elif request.workflow_type == "prd" and request.page_data:
            # PRD生成工作流也支持页面数据
            page_data = request.page_data
            
            # 构建页面上下文信息
            page_context = f"""
**当前页面信息**
- 页面标题: {page_data.get('title', '未知')}
- 页面域名: {page_data.get('domain', '未知')}
- 页面URL: {page_data.get('url', '未知')}"""
            
            # 添加技术栈信息
            tech_info = page_data.get('tech_info', {})
            if tech_info:
                page_context += f"""
**技术栈信息**
- 字体资源: {tech_info.get('fonts', [])}
- 样式表数量: {tech_info.get('stylesheets', 0)}
- 字体格式: {tech_info.get('font_formats', [])}"""
            
            # 添加页面内容摘要
            if page_data.get('text_summary'):
                page_context += f"""
**页面内容摘要**
{page_data.get('text_summary', '')[:300]}"""
            
            # 添加页面标题信息
            headings = page_data.get('headings', [])
            if headings:
                page_context += f"""
**页面主要标题**
{chr(10).join(headings[:5])}"""
            
            # 将页面上下文添加到用户查询中
            enhanced_query = f"{user_query}\n\n{page_context}"
            
            # 更新payload
            dify_payload["inputs"].update({
                "query": enhanced_query[:2000],  # 限制长度避免超限
                "text": enhanced_query[:2000],
                "page_domain": page_data.get('domain', '')[:50],
                "page_title": page_data.get('title', '')[:100],
            })
            
            print(f"添加页面数据到PRD工作流: {page_data.get('domain', 'unknown')}")
            print(f"增强查询长度: {len(enhanced_query)} 字符")
            print(f"传递给Dify的查询: {enhanced_query[:200]}...")
        
        
        # 按照经验：如果有图片，添加到inputs.input_img字段（必须是数组格式）
        if dify_files:
            # 只处理图片文件，其他类型文件暂时忽略
            image_files = [f for f in dify_files if f.get("type") == "image"]
            if image_files:
                dify_payload["inputs"]["input_img"] = image_files
                print(f"添加图片到input_img: {image_files}")
        
        print(f"最终Dify payload: {json.dumps(dify_payload, ensure_ascii=False, indent=2)}")

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        if request.stream:
            # 流式响应（SSE）
            print("处理流式响应请求（Dify）...")
            
            async def generate():
                full_content = ""  # 收集完整内容用于文件生成和数据库存储
                ai_message_id = None
                dify_conversation_id = conversation_id
                
                try:
                    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
                        # 根据官方文档：流式端点为 /chat-messages
                        async with client.stream(
                            "POST",
                            f"{DIFY_API_BASE}/chat-messages",
                            json=dify_payload,
                            headers={**headers, "Accept": "text/event-stream"},
                        ) as response:
                            if response.status_code != 200:
                                error_text = (await response.aread()).decode(errors="ignore")
                                error_msg = f"Dify API错误 (status={response.status_code}): {error_text}"
                                print(f"API错误: {error_msg}")
                                yield f"data: {json.dumps({"error": {"message": error_msg}}, ensure_ascii=False)}\n\n"
                                yield "data: [DONE]\n\n"
                                return

                            print("开始接收 Dify 流式数据...")
                            buffer = ""
                            current_task_id = None
                            async for chunk in response.aiter_text():
                                if chunk:
                                    buffer += chunk
                                    # 处理完整的行
                                    while '\n' in buffer:
                                        line, buffer = buffer.split('\n', 1)
                                        if line.strip():
                                            if line.startswith('data: '):
                                                # 解析 Dify 事件，转换为 OpenAI/DeepSeek 风格 delta
                                                try:
                                                    data = line[6:].strip()
                                                    if not data or data == '[DONE]':
                                                        pass
                                                    else:
                                                        parsed = json.loads(data)
                                                        event = parsed.get('event')

                                                        # 尽早把 task_id 传给前端，便于立即停止
                                                        if parsed.get('task_id') and parsed.get('task_id') != current_task_id:
                                                            current_task_id = parsed.get('task_id')
                                                            task_info = {"type": "task", "task_id": current_task_id}
                                                            yield f"data: {json.dumps(task_info, ensure_ascii=False)}\n\n"

                                                        if event == 'message':
                                                            # 追加内容
                                                            delta_text = parsed.get('answer', '')
                                                            if delta_text:
                                                                full_content += delta_text
                                                                transformed = {
                                                                    "choices": [
                                                                        {"delta": {"content": delta_text}}
                                                                    ]
                                                                }
                                                                # 若已知 task_id，一并携带
                                                                if current_task_id:
                                                                    transformed["task_id"] = current_task_id
                                                                yield f"data: {json.dumps(transformed, ensure_ascii=False)}\n\n"
                                                                # 同步流式渲染 HTML（忠实渲染，仅样式，不改语义）
                                                                try:
                                                                    html_partial = markdown.markdown(full_content, extensions=['extra', 'sane_lists'])
                                                                    html_evt = {"type": "html_partial", "html": html_partial}
                                                                    yield f"data: {json.dumps(html_evt, ensure_ascii=False)}\n\n"
                                                                except Exception as _e:
                                                                    print(f"流式 HTML 渲染失败: {_e}")
                                                            # 记录会话ID（如有）
                                                            if parsed.get('conversation_id'):
                                                                dify_conversation_id = parsed['conversation_id']
                                                            # 记录task_id（如有）
                                                            if current_task_id:
                                                                print(f"📋 从Dify获取到task_id: {current_task_id}")

                                                        elif event == 'message_file':
                                                            # 文件事件映射到前端可识别的自定义文件块
                                                            file_url = parsed.get('url')
                                                            file_type = parsed.get('type')
                                                            file_block = {
                                                                "type": "file",
                                                                "filename": f"{file_type}-{parsed.get('id')}",
                                                                "url": file_url,
                                                                "mime_type": "application/octet-stream",
                                                                "conversation_id": parsed.get('conversation_id') or dify_conversation_id,
                                                            }
                                                            yield f"data: {json.dumps(file_block, ensure_ascii=False)}\n\n"

                                                        elif event == 'message_end':
                                                            # 结束事件，保存会话ID
                                                            if parsed.get('conversation_id'):
                                                                dify_conversation_id = parsed['conversation_id']
                                                            # 不直接输出给前端（保持原有协议），将在结尾统一发 [DONE]

                                                        elif event in ('workflow_started', 'node_started', 'node_finished', 'workflow_finished', 'tts_message', 'tts_message_end', 'message_replace', 'ping'):
                                                            # 可选：忽略或在调试时打印
                                                            pass

                                                        elif event == 'error' or parsed.get('error'):
                                                            err_msg = parsed.get('message') or parsed.get('error') or '未知错误'
                                                            error_response = {"error": {"message": f"Dify 错误: {err_msg}"}}
                                                            yield f"data: {json.dumps(error_response, ensure_ascii=False)}\n\n"
                                                            yield "data: [DONE]\n\n"
                                                            return
                                                except Exception as e:
                                                    # 解析异常时忽略该行
                                                    print(f"解析Dify流式数据行失败: {e}")
                                                    pass
                            
                            # 处理剩余的buffer（无需额外转发）
                            if buffer.strip():
                                try:
                                    rem = buffer
                                    if rem.startswith('data: '):
                                        data = rem[6:].strip()
                                        if data and data != '[DONE]':
                                            parsed = json.loads(data)
                                            if parsed.get('event') == 'message':
                                                delta_text = parsed.get('answer', '')
                                                if delta_text:
                                                    full_content += delta_text
                                                    transformed = {"choices": [{"delta": {"content": delta_text}}]}
                                                    yield f"data: {json.dumps(transformed, ensure_ascii=False)}\n\n"
                                except Exception:
                                    pass
                            
                            # 将完整的AI回复存入数据库（始终使用本地会话ID）
                            if full_content:
                                ai_message_id = await db.add_message(
                                    conversation_id=conversation_id,
                                    role="assistant",
                                    content=full_content
                                )
                                print(f"添加AI回复到数据库: {ai_message_id}")

                            # 在结束前提供 Markdown 渲染后的 HTML（保持内容忠实，仅做渲染）
                            try:
                                if full_content and isinstance(full_content, str):
                                    html_output = markdown.markdown(full_content, extensions=['extra', 'sane_lists'])
                                    html_event = {"type": "html", "html": html_output}
                                    yield f"data: {json.dumps(html_event, ensure_ascii=False)}\n\n"
                            except Exception as _e:
                                print(f"HTML 渲染失败: {_e}")
                            
                            # 生成文件（如果需要）
                            if request.output_format != "text" and full_content.strip():
                                print(f"生成{request.output_format}文件...")
                                file_info = await generate_file(full_content, request.output_format)
                                if file_info:
                                    # 存储生成的文件信息到数据库
                                    if ai_message_id:
                                        file_db_id = await db.add_generated_file(
                                            message_id=ai_message_id,
                                            filename=file_info["filename"],
                                            file_path=os.path.join(GENERATED_DIR, file_info["filename"]),
                                            mime_type=file_info["mime_type"],
                                            format=request.output_format
                                        )
                                        print(f"添加生成的文件到数据库: {file_db_id}")
                                    
                                    # 向客户端发送文件信息
                                    file_data = {
                                        "type": "file",
                                        "filename": file_info["filename"],
                                        "url": file_info["url"],
                                        "mime_type": file_info["mime_type"],
                                        "conversation_id": dify_conversation_id or conversation_id
                                    }
                                    yield f"data: {json.dumps(file_data)}\n\n"
                            
                            # 向客户端发送对话ID（本地与 Dify）
                            conversation_data = {
                                "type": "conversation",
                                "conversation_id": conversation_id,  # 本地数据库会话ID
                                "dify_conversation_id": dify_conversation_id or ""
                            }
                            yield f"data: {json.dumps(conversation_data)}\n\n"
                            
                            yield "data: [DONE]\n\n"
                except httpx.StreamClosed:
                    print("流连接被关闭")
                    error_response = {
                        "error": {
                            "message": "Stream was closed unexpectedly"
                        }
                    }
                    yield f"data: {json.dumps(error_response)}\n\n"
                    yield "data: [DONE]\n\n"
                except Exception as e:
                    # 打印更完整的错误信息，便于定位
                    import traceback
                    print(f"流式处理错误: {repr(e)}")
                    print(traceback.format_exc())
                    error_response = {
                        "error": {
                            "message": f"流式处理错误: {repr(e)}"
                        }
                    }
                    yield f"data: {json.dumps(error_response)}\n\n"
                    yield "data: [DONE]\n\n"
            
            # 防止代理/浏览器缓冲：禁用缓存、保持连接
            return StreamingResponse(
                generate(),
                media_type="text/event-stream",
                headers={
                    "Cache-Control": "no-cache, no-transform",
                    "Pragma": "no-cache",
                    "Connection": "keep-alive",
                    "X-Accel-Buffering": "no",
                }
            )
        
        else:
            # 普通（阻塞）响应
            print("发送请求到 Dify API（阻塞模式）...")
            print(f"请求数据: {json.dumps(dify_payload, ensure_ascii=False, indent=2)}")
            
            async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
                response = await client.post(
                    _dify_endpoint(),
                    json=dify_payload,
                    headers=headers,
                )
                
                print(f"收到响应，状态码: {response.status_code}")
                if response.status_code != 200:
                    error_text = await response.aread()
                    error_text = error_text.decode('utf-8')
                    print(f"Dify API错误详情: {error_text}")
                    raise HTTPException(status_code=response.status_code, 
                                      detail=f"Dify API错误 (status={response.status_code}): {error_text}")
                
                dify_data = response.json()
                print(f"响应数据: {json.dumps(dify_data, ensure_ascii=False, indent=2)}")

                # Dify -> 前端兼容结构
                ai_content = dify_data.get('answer', '')
                # 存储AI回复到数据库（始终使用本地会话ID）
                if ai_content:
                    ai_message_id = await db.add_message(
                        conversation_id=conversation_id,
                        role="assistant",
                        content=ai_content
                    )
                    print(f"添加AI回复到数据库: {ai_message_id}")

                # 生成文件（如果需要）
                file_info = None
                if request.output_format != "text" and ai_content.strip():
                    print(f"生成{request.output_format}文件...")
                    file_info = await generate_file(ai_content, request.output_format)
                    if file_info and ai_content:
                        # 存储生成的文件信息到数据库
                        await db.add_generated_file(
                            message_id=ai_message_id,
                            filename=file_info["filename"],
                            file_path=os.path.join(GENERATED_DIR, file_info["filename"]),
                            mime_type=file_info["mime_type"],
                            format=request.output_format
                        )
                        print("生成文件已写入数据库")

                # 构造兼容的返回体
                # 读取 Dify 的会话ID用于前端持续对话
                dify_conversation_id_resp = dify_data.get('conversation_id') or ""

                compatible = {
                    "id": dify_data.get('id', str(uuid.uuid4())),
                    "object": "chat.completion",
                    "created": int(datetime.now().timestamp()),
                    "model": request.model,
                    "choices": [
                        {
                            "index": 0,
                            "message": {"role": "assistant", "content": ai_content},
                            "finish_reason": "stop",
                        }
                    ],
                    "usage": dify_data.get('metadata', {}).get('usage', {
                        "prompt_tokens": 0,
                        "completion_tokens": 0,
                        "total_tokens": 0,
                    }),
                    # 对前端：conversation_id 为本地会话ID；另返回 dify_conversation_id
                    "conversation_id": conversation_id,
                    "dify_conversation_id": dify_conversation_id_resp,
                }

                # 提供 HTML 渲染字段（保持内容忠实，仅渲染样式）
                try:
                    if ai_content and isinstance(ai_content, str):
                        compatible['html'] = markdown.markdown(ai_content, extensions=['extra', 'sane_lists'])
                except Exception as _e:
                    print(f"阻塞模式 HTML 渲染失败: {_e}")

                if file_info:
                    compatible['file'] = file_info

                return compatible
                
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="请求超时")
    except httpx.RequestError as e:
        print(f"网络请求错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"网络请求错误: {str(e)}")
    except Exception as e:
        print(f"服务器内部错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {str(e)}")

@app.post("/api/chat/stop/{task_id}")
async def stop_chat_response(task_id: str, api_key: str = "prd"):
    """
    停止指定的Dify响应任务
    api_key参数指定使用哪个API Key ("prd" 或 "prompt")
    """
    
    # 根据api_key参数选择对应的API Key
    if api_key == "prompt":
        dify_api_key = DIFY_PROMPT_API_KEY
        if not dify_api_key:
            raise HTTPException(status_code=500, detail="Prompt生成工作流API密钥未配置")
    else:  # 默认使用PRD工作流
        dify_api_key = DIFY_PRD_API_KEY
        if not dify_api_key:
            raise HTTPException(status_code=500, detail="PRD生成工作流API密钥未配置")
    
    try:
        print(f"停止响应任务: {task_id}")
        
        # 调用Dify API停止响应
        headers = {
            "Authorization": f"Bearer {dify_api_key}",
            "Content-Type": "application/json"
        }
        
        # 构造候选停止端点（不同网关/版本可能不同）
        # 为了兼容性，不再区分 channel，全部都尝试
        base_candidates = [
            f"{DIFY_API_BASE}/chat-messages/{task_id}/stop",
            f"{DIFY_API_BASE}/workflows/{task_id}/stop",
            f"{DIFY_API_BASE}/workflows/tasks/{task_id}/stop",
            f"{DIFY_API_BASE}/tasks/{task_id}/stop",
        ]
        # 同时加入带斜杠版本，解决 308 重定向（有些网关要求以 / 结尾）
        candidate_urls = []
        for u in base_candidates:
            candidate_urls.append(u)
            if not u.endswith('/'):
                candidate_urls.append(u + '/')

        last_status = None
        last_text = None
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            for stop_url in candidate_urls:
                try:
                    print(f"尝试停止 URL: {stop_url}")
                    user_val = "abc-123"
                    # 按文档先用 JSON body
                    response = await client.post(stop_url, headers={**headers, "Accept": "application/json"}, json={"user": user_val})
                    last_status = response.status_code
                    last_text = response.text
                    print(f"停止返回: {last_status} {last_text[:200]}")
                    if response.status_code in (200, 204):
                        try:
                            dify_response = response.json()
                        except Exception:
                            dify_response = {"result": "success"}
                        return {
                            "success": True,
                            "result": dify_response.get("result", "success"),
                            "message": "响应已成功停止"
                        }

                    # 再尝试 Query 方式 ?user=xxx（部分网关可能这样要求）
                    response = await client.post(stop_url, headers=headers, params={"user": user_val})
                    last_status = response.status_code
                    last_text = response.text
                    print(f"停止返回(带query): {last_status} {last_text[:200]}")
                    if response.status_code in (200, 204):
                        try:
                            dify_response = response.json()
                        except Exception:
                            dify_response = {"result": "success"}
                        return {
                            "success": True,
                            "result": dify_response.get("result", "success"),
                            "message": "响应已成功停止"
                        }

                    # 再尝试自定义 Header 方式 X-User（防御性兼容）
                    response = await client.post(stop_url, headers={**headers, "X-User": user_val})
                    last_status = response.status_code
                    last_text = response.text
                    print(f"停止返回(带X-User): {last_status} {last_text[:200]}")
                    if response.status_code in (200, 204):
                        try:
                            dify_response = response.json()
                        except Exception:
                            dify_response = {"result": "success"}
                        return {
                            "success": True,
                            "result": dify_response.get("result", "success"),
                            "message": "响应已成功停止"
                        }
                except Exception as e:
                    print(f"调用 {stop_url} 出错: {e}")

        raise HTTPException(
            status_code=500 if (last_status is None or last_status == 200) else last_status,
            detail=f"停止响应失败: {last_text or '未知错误'}"
        )
                
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="停止响应请求超时")
    except httpx.RequestError as e:
        print(f"停止响应网络错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"网络错误: {str(e)}")
    except Exception as e:
        print(f"停止响应内部错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"内部错误: {str(e)}")

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), message_id: Optional[str] = Form(None), conversation_id: Optional[str] = Form(None)):
    """
    上传文件
    
    如果提供了message_id，则将文件关联到指定消息
    如果提供了conversation_id但没有message_id，则创建一个新的系统消息并关联文件
    """
    try:
        # 验证文件类型
        allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png', '.gif', '.xlsx', '.xls', '.pptx', '.ppt'}
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(status_code=400, detail="不支持的文件类型")
        
        # 验证文件大小 (50MB)
        content = await file.read()
        file_size = len(content)
        if file_size > 50 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="文件大小超过50MB限制")
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        file_extension = os.path.splitext(file.filename)[1]
        new_filename = f"{file_id}{file_extension}"
        file_path = os.path.join(UPLOAD_DIR, new_filename)
        
        # 保存文件
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)
        
        # 处理数据库存储
        attachment_id = None
        
        if message_id:
            # 将文件关联到现有消息
            attachment_id = await db.add_attachment(
                message_id=message_id,
                filename=file.filename,
                file_path=file_path,
                mime_type=file.content_type
            )
            print(f"文件 {file.filename} 关联到消息 {message_id}, 附件ID: {attachment_id}")
            
        elif conversation_id:
            # 验证对话是否存在
            conversation = await db.get_conversation(conversation_id)
            if not conversation:
                raise HTTPException(status_code=404, detail="指定的对话不存在")
            
            # 创建一个系统消息并关联文件
            system_message_id = await db.add_message(
                conversation_id=conversation_id,
                role="system",
                content=f"上传了文件: {file.filename}"
            )
            
            attachment_id = await db.add_attachment(
                message_id=system_message_id,
                filename=file.filename,
                file_path=file_path,
                mime_type=file.content_type
            )
            print(f"文件 {file.filename} 关联到新的系统消息 {system_message_id}, 附件ID: {attachment_id}")
        
        return {
            "file_id": file_id,
            "filename": file.filename,
            "size": file_size,
            "type": file.content_type,
            "uploaded_at": datetime.now().isoformat(),
            "attachment_id": attachment_id,
            "message_id": message_id,
            "conversation_id": conversation_id
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件上传失败: {str(e)}")

@app.get("/api/files/{filename}")
async def download_file(filename: str):
    """
    下载生成的文件
    """
    file_path = os.path.join(GENERATED_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 根据文件扩展名设置MIME类型
    if filename.endswith('.pdf'):
        media_type = 'application/pdf'
    elif filename.endswith('.docx'):
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif filename.endswith('.md'):
        media_type = 'text/markdown'
    else:
        media_type = 'application/octet-stream'
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=media_type
    )

@app.get("/api/models")
async def get_available_models():
    """
    获取可用的模型列表
    """
    return {
        "models": [
            {
                "id": "deepseek-chat",
                "name": "DeepSeek Chat",
                "description": "DeepSeek通用对话模型",
                "max_tokens": 8192
            },
            {
                "id": "deepseek-coder",
                "name": "DeepSeek Coder", 
                "description": "DeepSeek代码生成模型",
                "max_tokens": 8192
            }
        ]
    }

@app.get("/api/conversations")
async def get_conversations(limit: int = 20, offset: int = 0):
    """获取对话列表"""
    try:
        conversations = await db.get_conversations(limit=limit, offset=offset)
        return {"conversations": conversations}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取对话列表失败: {str(e)}")

@app.get("/api/conversations/{conversation_id}")
async def get_conversation(conversation_id: str):
    """获取对话详情"""
    try:
        conversation = await db.get_conversation(conversation_id)
        if not conversation:
            raise HTTPException(status_code=404, detail="对话不存在")
        
        messages = await db.get_messages(conversation_id)
        return {
            "conversation": conversation,
            "messages": messages
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取对话详情失败: {str(e)}")

@app.put("/api/conversations/{conversation_id}")
async def update_conversation(conversation_id: str, title: str):
    """更新对话标题"""
    try:
        success = await db.update_conversation_title(conversation_id, title)
        if not success:
            raise HTTPException(status_code=404, detail="对话不存在")
        return {"status": "success"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"更新对话失败: {str(e)}")

@app.delete("/api/conversations/{conversation_id}")
async def delete_conversation(conversation_id: str):
    """删除对话"""
    try:
        success = await db.delete_conversation(conversation_id)
        if not success:
            raise HTTPException(status_code=404, detail="对话不存在")
        return {"status": "success"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除对话失败: {str(e)}")

@app.get("/api/stats")
async def get_stats():
    """获取统计信息"""
    try:
        stats = await db.get_stats()
        return stats
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取统计信息失败: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)